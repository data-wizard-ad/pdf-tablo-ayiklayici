import streamlit as st
import pdfplumber
import pandas as pd
from io import BytesIO
import re
from PIL import Image
import numpy as np
import json
from pypdf import PdfReader, PdfWriter 

# --- GÜVENLİ WORD İTHALATI ---
try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# --- 1. GÜVENLİ OCR İTHALATI ---
try:
    import easyocr
    @st.cache_resource
    def load_ocr(): return easyocr.Reader(['tr', 'en'])
    reader = load_ocr()
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# --- 2. SEO VE SAYFA AYARLARI ---
st.set_page_config(
    page_title="Master Veri Sihirbazı Elite | AI Özetleme & PDF OCR",
    page_icon="🪄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- YARDIMCI FONKSİYONLAR ---

def get_pdf_preview(pdf_file, page_no=0):
    """Belirli bir PDF sayfasını ön izleme için görsele çevirir."""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            if page_no < len(pdf.pages):
                page = pdf.pages[page_no]
                return page.to_image(resolution=72).original
    except:
        return None
    return None

def to_word(df):
    if not WORD_AVAILABLE: return None
    doc = Document()
    doc.add_heading('Data Wizard Elite Veri Raporu', 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def pdf_to_word_direct(pdf_file):
    if not WORD_AVAILABLE: return None
    doc = Document()
    doc.add_heading('PDF Metin Aktarımı', 0)
    reader_word = PdfReader(pdf_file)
    for page in reader_word.pages:
        doc.add_paragraph(page.extract_text())
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- MANİPÜLASYON FONKSİYONLARI ---

def add_page_numbers(input_pdf):
    from reportlab.pdfgen import canvas
    reader_num = PdfReader(input_pdf)
    writer_num = PdfWriter()
    for i in range(len(reader_num.pages)):
        page = reader_num.pages[i]
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=(width, height))
        can.setFont("Helvetica", 10)
        can.drawString(width - 50, 30, f"{i + 1}")
        can.save()
        packet.seek(0)
        num_pdf = PdfReader(packet)
        page.merge_page(num_pdf.pages[0])
        writer_num.add_page(page)
    bio = BytesIO()
    writer_num.write(bio)
    return bio.getvalue()

def compress_pdf(input_pdf):
    reader_comp = PdfReader(input_pdf)
    writer_comp = PdfWriter()
    for page in reader_comp.pages:
        new_page = writer_comp.add_page(page)
        new_page.compress_content_streams()
    bio = BytesIO()
    writer_comp.write(bio)
    return bio.getvalue()

def rotate_pdf(input_pdf, rotation_angle):
    reader_rot = PdfReader(input_pdf)
    writer_rot = PdfWriter()
    for page in reader_rot.pages:
        page.rotate(rotation_angle)
        writer_rot.add_page(page)
    bio = BytesIO()
    writer_rot.write(bio)
    return bio.getvalue()

def split_pdf(input_pdf, start_page, end_page):
    reader_split = PdfReader(input_pdf)
    writer_split = PdfWriter()
    for i in range(start_page - 1, min(end_page, len(reader_split.pages))):
        writer_split.add_page(reader_split.pages[i])
    bio = BytesIO()
    writer_split.write(bio)
    return bio.getvalue()

def encrypt_pdf(input_pdf, password):
    reader_enc = PdfReader(input_pdf)
    writer_enc = PdfWriter()
    for page in reader_enc.pages:
        writer_enc.add_page(page)
    writer_enc.encrypt(password)
    bio = BytesIO()
    writer_enc.write(bio)
    return bio.getvalue()

def images_to_pdf(image_files):
    img_list = []
    for img_file in image_files:
        img = Image.open(img_file)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img_list.append(img)
    if img_list:
        bio = BytesIO()
        img_list[0].save(bio, format="PDF", save_all=True, append_images=img_list[1:])
        return bio.getvalue()
    return None

def convert_image(img_file, target_format):
    img = Image.open(img_file)
    if target_format.upper() in ["JPG", "JPEG"] and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    out_img = BytesIO()
    if target_format.upper() == "ICO":
        img.save(out_img, format="ICO", sizes=[(256, 256), (128, 128), (64, 64), (32, 32)])
    else:
        img.save(out_img, format=target_format.upper())
    return out_img.getvalue()

# SEO ve Google Analiz
st.markdown("""<script async src="https://www.googletagmanager.com/gtag/js?id=G-SH8W61QFSS"></script>
<script>
  window.dataLayer = window.dataLayer || [];
  function gtag(){dataLayer.push(arguments);}
  gtag('js', new Date());
  gtag('config', 'G-SH8W61QFSS');
</script>""", unsafe_allow_html=True)

# --- 3. YAN MENÜ (SİDEBAR) ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3652/3652191.png", width=70)
    st.title("Wizard Global")
    st.divider()
    pdf_password = st.text_input("🔑 PDF Şifresi (Varsa)", type="password")
    lang = st.selectbox("🌐 Dil / Language", ["Türkçe", "English"], index=0)
    st.info("🛡️ Verileriniz yerel RAM'de işlenir.")
    st.divider()
    ai_insights = st.toggle("Yapay Zeka Analizi & Özet", value=True)
    show_charts = st.toggle("Grafik Analizini Göster", value=True)
    st.divider()
    st.link_button("☕ Kahve Ismarla", "https://buymeacoffee.com/databpak", use_container_width=True)
    st.caption("v4.5.0 PREVIEW | BY BERKANT PAK | 2026")

# --- 4. ÜST BİLGİ KARTLARI ---
col1, col2, col3, col4 = st.columns(4)
with col1: st.metric("İşleme", "Yerel (Edge)", "Encrypted")
with col2: st.metric("Güvenlik", "Shield Active", "Shielded")
with col3: st.metric("Etki", "22+ Kullanıcı", "Growing")
with col4: st.metric("Lisans", "Open-Source", "MIT")
st.divider()

# --- 5. ANA PANEL ---
st.title("🧙‍♂️ Master Veri Sihirbazı Elite")

tab1, tab2, tab3 = st.tabs(["📄 PDF İşleme", "🖼️ Resimden Yazıya (OCR)", "🛠️ Editör & Dönüştürücü"])

with tab1:
    pdf_files = st.file_uploader("PDF dosyalarını yükleyin", type="pdf", accept_multiple_files=True, key="pdf_table_uploader")
    if pdf_files:
        all_data = {}
        with st.status("🔮 Sihirbaz dosyaları inceliyor...", expanded=False) as status:
            for f in pdf_files:
                try:
                    img_preview = get_pdf_preview(f)
                    if img_preview:
                        st.image(img_preview, caption=f"{f.name} - İlk Sayfa Ön İzlemesi", width=300)
                    with pdfplumber.open(f, password=pdf_password) as pdf:
                        pages_list = []
                        for i, page in enumerate(pdf.pages):
                            table = page.extract_table()
                            if table:
                                df = pd.DataFrame(table[1:], columns=table[0])
                                df.columns = [f"Kol_{idx}" if not c else str(c) for idx, c in enumerate(df.columns)]
                                pages_list.append((f"Sayfa {i+1}", df))
                        if pages_list: all_data[f.name] = pages_list
                except Exception as e:
                    st.error(f"⚠️ Hata: {str(e)}")
            if all_data: status.update(label="✅ İşlem Tamam!", state="complete")

        if all_data:
            sel_file = st.selectbox("Dosya seçin:", list(all_data.keys()))
            file_pages = all_data.get(sel_file, [])
            if file_pages:
                pdf_tabs = st.tabs([t[0] for t in file_pages])
                for i, (p_name, df) in enumerate(file_pages):
                    with pdf_tabs[i]:
                        st.dataframe(df, use_container_width=True)
                        st.divider()
                        d_col1, d_col2, d_col3 = st.columns(3)
                        with d_col1:
                            out_ex = BytesIO(); df.to_excel(out_ex, index=False)
                            st.download_button("📂 Excel İndir", out_ex.getvalue(), f"{p_name}.xlsx", key=f"ex_v_{i}")
                        with d_col2:
                            st.download_button("📄 CSV İndir", df.to_csv(index=False).encode('utf-8-sig'), f"{p_name}.csv", key=f"csv_v_{i}")
                        with d_col3:
                            word_data = to_word(df)
                            if word_data: st.download_button("📝 Word İndir", word_data, f"{p_name}.docx", key=f"word_v_{i}")

with tab2:
    st.subheader("🖼️ Görselden Veri Ayıklama")
    uploaded_img = st.file_uploader("Resim yükleyin", type=["jpg", "png", "jpeg"])
    if uploaded_img:
        st.image(uploaded_img, caption="Yüklenen Görsel", width=400)
        if st.button("🚀 Resmi Tara"):
            if OCR_AVAILABLE:
                with st.spinner("🧠 Metinler okunuyor..."):
                    result = reader.readtext(np.array(Image.open(uploaded_img)), detail=0)
                    st.table(pd.DataFrame(result, columns=["Ayıklanan Veriler"]))

with tab3:
    col_tools, col_conv = st.columns([1, 1])
    
    with col_tools:
        st.subheader("🛠️ PDF Araçları")
        edit_mode = st.selectbox("İşlem Seçin:", [
            "PDF Birleştirme", "Sayfa Ayırma", "PDF Sayfalarını Döndür",
            "🔢 Sayfa Numarası Ekle",
            "🚫 Filigran Kaldır Pro",
            "🔄 Sayfa Sıralamasını Değiştir",
            "🗑️ Sayfa Sil / Sırala",
            "🔐 PDF Şifrele (Parola Koy)", "🖼️ Görsellerden PDF Yap",
            "PDF to Word (Direkt)", "📉 PDF Boyutu Küçült"
        ])
        
        preview_container = st.empty()

        if edit_mode == "PDF Birleştirme":
            merge_files = st.file_uploader("Birleştirilecek PDF'ler", type="pdf", accept_multiple_files=True, key="m_up_fix")
            if merge_files:
                st.write(f"📂 {len(merge_files)} dosya hazır.")
                if st.button("🔗 Birleştir"):
                    merger = PdfWriter()
                    for pdf in merge_files: merger.append(pdf)
                    out = BytesIO(); merger.write(out)
                    st.download_button("📥 İndir", out.getvalue(), "birlesmis.pdf")
      
        elif edit_mode == "🔢 Sayfa Numarası Ekle":
            num_file = st.file_uploader("Numara eklenecek PDF", type="pdf", key="num_up")
            if num_file:
                img = get_pdf_preview(num_file)
                if img: preview_container.image(img, caption="İşlem Öncesi Görünüm", width=250)
                if st.button("🔢 Numaraları Bas ve Hazırla"):
                    try:
                        with st.spinner("Sihirbaz sayfaları mühürlüyor..."):
                            numbered_pdf = add_page_numbers(num_file)
                            st.success("✅ Tüm sayfalar numaralandırıldı!")
                            st.download_button("📥 Numaralı PDF'i İndir", numbered_pdf, "wizard_numbered.pdf")
                    except Exception as e:
                        st.error(f"Hata: {e}")

        elif edit_mode == "🚫 Filigran Kaldır Pro":
            wm_file = st.file_uploader("Filigranlı PDF seçin", type="pdf", key="wm_pro_v5_final")
            if wm_file:
                img = get_pdf_preview(wm_file)
                if img: st.image(img, width=250, caption="İşlem Yapılacak Dosya")
                
                clean_method = st.radio("Yöntem:", ["Hassas Metin Temizliği", "Nokta Atışı Maskeleme (Önerilen)"])
                
                if clean_method == "Hassas Metin Temizliği":
                    target_text = st.text_input("Silinecek Metin", "iLovePDF")
                    if st.button("🧼 Metni Kazı"):
                        reader_wm = PdfReader(wm_file)
                        writer_wm = PdfWriter()
                        for page in reader_wm.pages:
                            new_page = writer_wm.add_page(page)
                            contents = new_page.get_contents()
                            if contents:
                                data = contents.get_data() if not isinstance(contents, list) else b"".join([c.get_data() for c in contents])
                                if target_text.encode() in data:
                                    new_page.get_contents().set_data(data.replace(target_text.encode(), b" "))
                        out = BytesIO(); writer_wm.write(out)
                        st.download_button("📥 İndir", out.getvalue(), "metin_temiz.pdf")

                else:
                    st.info("💡 iLovePDF filigranları için optimize edilmiş hassas maskeleme uygulanır.")
                    if st.button("⬜ Hassas Maskeleme Başlat"):
                        from reportlab.pdfgen import canvas
                        from reportlab.lib.colors import white
                        reader_mask = PdfReader(wm_file)
                        writer_mask = PdfWriter()
                        
                        for page in reader_mask.pages:
                            w, h = float(page.mediabox.width), float(page.mediabox.height)
                            packet = BytesIO()
                            can = canvas.Canvas(packet, pagesize=(w, h))
                            can.setFillColor(white)
                            can.setStrokeColor(white)
                            
                            # iLovePDF Dağılım Koordinatları (Hassaslaştırılmış Küçük Kutular)
                            # Köşeler
                            can.rect(5, 5, 60, 15, fill=1) # Sol Alt
                            can.rect(w-65, 5, 60, 15, fill=1) # Sağ Alt
                            can.rect(5, h-20, 60, 15, fill=1) # Sol Üst
                            can.rect(w-65, h-20, 60, 15, fill=1) # Sağ Üst
                            
                            # Orta ve Kenar Dağılımları (Metni bozmamak için çok dar tutuldu)
                            can.rect(w/2 - 30, 5, 60, 15, fill=1) # Alt Orta
                            can.rect(w/2 - 30, h-20, 60, 15, fill=1) # Üst Orta
                            can.rect(5, h/2 - 10, 60, 15, fill=1) # Sol Orta
                            can.rect(w-65, h/2 - 10, 60, 15, fill=1) # Sağ Orta
                            
                            can.save()
                            packet.seek(0)
                            page.merge_page(PdfReader(packet).pages[0])
                            writer_mask.add_page(page)
                            
                        out = BytesIO(); writer_mask.write(out)
                        st.success("✅ Metne zarar vermeden kenar filigranları maskelendi.")
                        st.download_button("📥 Temiz PDF İndir", out.getvalue(), "hassas_maske.pdf")
        elif edit_mode == "🔄 Sayfa Sıralamasını Değiştir":
            reorder_file = st.file_uploader("PDF seçin", type="pdf", key="reorder_up")
            if reorder_file:
                reader_re = PdfReader(reorder_file)
                total_p = len(reader_re.pages)
                page_indices = list(range(total_p))
                new_order_indices = st.multiselect("Yeni Sıra:", options=page_indices, default=page_indices, format_func=lambda x: f"Sayfa {x + 1}")
                if st.button("🪄 Yeni Sırayla Oluştur") and new_order_indices:
                    writer_re = PdfWriter()
                    for p_idx in new_order_indices: writer_re.add_page(reader_re.pages[p_idx])
                    out = BytesIO(); writer_re.write(out)
                    st.download_button("📥 İndir", out.getvalue(), "reordered.pdf")

        elif edit_mode == "🗑️ Sayfa Sil / Sırala":
            sort_file = st.file_uploader("PDF seçin", type="pdf", key="sort_unique_key")
            if sort_file:
                reader_sort = PdfReader(sort_file)
                selected_indices = st.multiselect("Tutulacak Sayfalar:", options=list(range(len(reader_sort.pages))), default=list(range(len(reader_sort.pages))), format_func=lambda x: f"Sayfa {x+1}")
                if st.button("🪄 Yeni PDF Oluştur"):
                    writer_sort = PdfWriter()
                    for idx in selected_indices: writer_sort.add_page(reader_sort.pages[idx])
                    out = BytesIO(); writer_sort.write(out)
                    st.download_button("📥 İndir", out.getvalue(), "edited.pdf")

        elif edit_mode == "Sayfa Ayırma":
            split_file = st.file_uploader("PDF seçin", type="pdf", key="sp_up")
            if split_file:
                reader_sp = PdfReader(split_file)
                c1, c2 = st.columns(2)
                start_p = c1.number_input("Başlangıç", min_value=1, max_value=len(reader_sp.pages), value=1)
                end_p = c2.number_input("Bitiş", min_value=1, max_value=len(reader_sp.pages), value=len(reader_sp.pages))
                if st.button("✂️ Kes ve Ayır"):
                    st.download_button("📥 İndir", split_pdf(split_file, start_p, end_p), "split.pdf")

        elif edit_mode == "PDF Sayfalarını Döndür":
            rot_file = st.file_uploader("PDF seçin", type="pdf", key="rot_up")
            if rot_file:
                angle = st.radio("Açı:", [90, 180, 270], horizontal=True)
                if st.button("🔄 Döndür"):
                    st.download_button("📥 İndir", rotate_pdf(rot_file, angle), "rotated.pdf")

        elif edit_mode == "🔐 PDF Şifrele (Parola Koy)":
            enc_file = st.file_uploader("Şifrelenecek PDF", type="pdf", key="enc_up")
            if enc_file:
                new_pass = st.text_input("Şifre", type="password")
                if st.button("🔒 Şifrele") and new_pass:
                    st.download_button("📥 İndir", encrypt_pdf(enc_file, new_pass), "encrypted.pdf")

        elif edit_mode == "🖼️ Görsellerden PDF Yap":
            port_files = st.file_uploader("Resimleri Seçin", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
            if port_files and st.button("📑 PDF Yap"):
                st.download_button("📥 İndir", images_to_pdf(port_files), "images_to.pdf")

        elif edit_mode == "PDF to Word (Direkt)":
            word_file = st.file_uploader("PDF seçin", type="pdf", key="word_direct_up")
            if word_file and st.button("📝 Dönüştür"):
                st.download_button("📥 İndir", pdf_to_word_direct(word_file), "converted.docx")

        elif edit_mode == "📉 PDF Boyutu Küçült":
            comp_file = st.file_uploader("Küçültülecek PDF", type="pdf", key="comp_up")
            if comp_file and st.button("🚀 Optimize Et"):
                st.download_button("📥 İndir", compress_pdf(comp_file), "compressed.pdf")

    with col_conv:
        st.subheader("🔄 Görsel Dönüştürücü")
        img_conv_file = st.file_uploader("Görsel yükleyin", type=["jpg", "jpeg", "png", "webp", "bmp"], key="img_conv_final_unique")
        if img_conv_file:
            st.image(img_conv_file, width=150)
            target_ext = st.selectbox("Hedef Format:", ["PNG", "JPG", "ICO", "WEBP", "BMP"])
            if st.button(f"✨ Dönüştür"):
                converted_bytes = convert_image(img_conv_file, target_ext)
                st.download_button(f"📥 {target_ext} İndir", converted_bytes, f"conv.{target_ext.lower()}")


